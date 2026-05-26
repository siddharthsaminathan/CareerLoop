"""
CV file extraction utilities.
Converts PDF and DOCX binary content to plain text for LLM processing.
"""

import io
import logging
from typing import Optional

logger = logging.getLogger("careerloop.onboarding.cv_extractor")


def extract_pdf_text(content: bytes) -> Optional[str]:
    """Extract plain text from PDF bytes. Tries pypdf first, falls back to pdfminer."""
    # Try pypdf (lighter, faster)
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(pages).strip()
        if text:
            return text
    except Exception as e:
        logger.debug("pypdf failed: %s", e)

    # Fallback: pdfminer.six (more accurate for complex layouts)
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(io.BytesIO(content))
        if text and text.strip():
            return text.strip()
    except Exception as e:
        logger.debug("pdfminer failed: %s", e)

    logger.warning("All PDF extractors failed — no text extracted")
    return None


def extract_docx_text(content: bytes) -> Optional[str]:
    """Extract plain text from DOCX bytes using python-docx."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table cell text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs) if paragraphs else None
    except Exception as e:
        logger.error("DOCX extraction failed: %s", e)
        return None


def extract_text_from_bytes(content: bytes, mime_type: str = "") -> Optional[str]:
    """
    Auto-detect format from mime_type and extract text.
    Supports: application/pdf, application/vnd.openxmlformats-officedocument.wordprocessingml.document
    """
    mime = mime_type.lower()
    if "pdf" in mime:
        return extract_pdf_text(content)
    if "word" in mime or "docx" in mime or "openxmlformats" in mime:
        return extract_docx_text(content)
    # Plain text fallback
    try:
        return content.decode("utf-8", errors="replace").strip() or None
    except Exception:
        return None
